"""RFI / 招標文件合規檢核（Story 006 Phase 3）。

讀使用者上傳的 RFI / 需求書文字，用正規表示式比對採購法條文常見違規訊號。
這是 heuristics + 啟發式提示，**不是法律結論**。
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass


@dataclass
class DocFinding:
    rule_id: str
    article_no: str
    article_title: str
    severity: str  # low / mid / high
    snippet: str
    advice: str


# ---------- 文字抽取 ----------


def extract_text(filename: str, data: bytes) -> tuple[str, str]:
    """回傳 (text, note)。note 為抽取狀態（成功/警告/退化）。"""
    name = filename.lower()
    if name.endswith(".txt") or name.endswith(".md"):
        for enc in ("utf-8", "big5", "cp950"):
            try:
                return data.decode(enc), f"以 {enc} 解碼"
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace"), "編碼不明，已強制 UTF-8"

    if name.endswith(".pdf"):
        try:
            import pypdf
        except ImportError:
            return "", "缺少 pypdf 套件，無法解析 PDF（pip install pypdf）"
        try:
            reader = pypdf.PdfReader(io.BytesIO(data))
            pages = [p.extract_text() or "" for p in reader.pages]
            txt = "\n".join(pages)
            note = f"PDF {len(reader.pages)} 頁，抽出 {len(txt):,} 字"
            if len(txt.strip()) < 50:
                note += "（疑似掃描檔，需 OCR）"
            return txt, note
        except Exception as e:
            return "", f"PDF 解析失敗：{e}"

    if name.endswith(".docx"):
        try:
            import docx  # python-docx
        except ImportError:
            return "", "缺少 python-docx 套件（pip install python-docx）"
        try:
            d = docx.Document(io.BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.extend(c.text for c in row.cells)
            txt = "\n".join(parts)
            return txt, f"DOCX 段落 {len(d.paragraphs)}，抽出 {len(txt):,} 字"
        except Exception as e:
            return "", f"DOCX 解析失敗：{e}"

    return "", f"不支援的副檔名：{filename}"


# ---------- 規則 ----------

# 中文數字轉阿拉伯（給「壹億」「一億」這類數字偵測）
CN_DIGIT = {"零": 0, "一": 1, "壹": 1, "二": 2, "貳": 2, "兩": 2, "三": 3, "參": 3,
            "四": 4, "肆": 4, "五": 5, "伍": 5, "六": 6, "陸": 6, "七": 7, "柒": 7,
            "八": 8, "捌": 8, "九": 9, "玖": 9, "十": 10, "拾": 10}


def _ctx(text: str, start: int, end: int, pad: int = 40) -> str:
    a = max(0, start - pad)
    b = min(len(text), end + pad)
    return text[a:b].replace("\n", " ")


# 出現在 match 前 ~15 字內 → 視為「否定／引用法條／教育性敘述」，不算違規
_NEGATION_PREFIX_RE = re.compile(
    r"(不得|不可|禁止|非採|未採|非屬|非以|並非|不採|依法|依規定|依.{0,4}標準|"
    r"採購法.{0,8}?條|本法.{0,4}?規定|得標公告|定義|認定標準|名詞解釋)"
)


def _is_negated_or_quoted(text: str, start: int, lookback: int = 20) -> bool:
    win = text[max(0, start - lookback):start]
    return bool(_NEGATION_PREFIX_RE.search(win))


# 「不得少於／不得低於／不得短於 N 日」「未滿 N 日以 N 日計」「原等標期之 N 分之 N」
# — 引用法定下限／計算規則的句型，後面那個數字不代表本案實際期限
_LOWER_BOUND_QUOTE_RE = re.compile(
    r"不得\s*(少|低|短)\s*於|至少|最少|未滿|未達|"
    r"原(等標期|履約期限|期限).{0,6}?分之|"
    r"分之[一二三四五六七八九十\d]+"
)


def _is_lower_bound_quote(text: str, num_start: int, lookback: int = 25) -> bool:
    win = text[max(0, num_start - lookback):num_start]
    return bool(_LOWER_BOUND_QUOTE_RE.search(win))


# 帶有強烈 brand spec 訊號的 label 與 value pattern
# 規則：label 後緊接英文/數字字串（≥2 字、含至少一個字母或數字），才視為「指定特定廠牌」。
# 純中文後綴（如「廠牌型號」「原廠物料」「新型號案件」）一律不算。
_BRAND_VALUE = r"[A-Za-z][A-Za-z0-9\-\s\.]{1,30}|\d{2,}[A-Za-z0-9\-]*"
_BRAND_LABEL_VALUE_RE = re.compile(
    r"(廠牌|品牌|指定型號|型號|機型)\s*[:：]\s*(" + _BRAND_VALUE + r")"
    r"|(廠牌|品牌|型號|機型)\s*[，、]?\s*(" + _BRAND_VALUE + r")"
)

# 屬於常見「合法或無關」的廠牌相關片語，直接整段忽略
_BRAND_FALSE_POSITIVE_RE = re.compile(
    r"原廠\s*(物料|維修|保固|支援|服務|授權|技術|設備|耗材|零件|料件)"
    r"|(中國大陸|大陸|中國)\s*廠牌"
    r"|新\s*型號|舊\s*型號|型號\s*變更|型號\s*規格|廠牌\s*型號|廠牌\s*資料|廠牌\s*名稱|廠牌\s*清單"
)

_BLANKET_EQUIV_RE = re.compile(
    r"(本(文件|案|招標文件)|所列|規格).{0,30}?"
    r"(廠牌|品牌|型號|規格).{0,30}?(均得|皆得|可|得).{0,10}?(以|用)?\s*同等品"
    r"|本(文件|案|招標文件).{0,40}?同等品.{0,10}?(代替|替代|皆可|均可)"
)


def check_brand_without_equivalent(text: str) -> list[DocFinding]:
    """§26：規格不得指定特定品牌，除非註明『或同等品』。

    若文件有「總則式同等品聲明」（例：本文件所列廠牌型號均得以同等品代替），
    視為已符合 §26 第 3 項，整份不報。
    """
    if _BLANKET_EQUIV_RE.search(text):
        return []
    out = []
    for m in _BRAND_LABEL_VALUE_RE.finditer(text):
        if _is_negated_or_quoted(text, m.start()):
            continue
        # 命中位置前後 8 字若為常見合法/無關片語（原廠物料、中國大陸廠牌、新型號…）→ skip
        win = text[max(0, m.start() - 8):min(len(text), m.end() + 8)]
        if _BRAND_FALSE_POSITIVE_RE.search(win):
            continue
        # 同行有「或同等品」即視為合規
        line_start = text.rfind("\n", 0, m.start()) + 1
        line_end = text.find("\n", m.end())
        if line_end == -1:
            line_end = len(text)
        line = text[line_start:line_end]
        if re.search(r"或同等品|等同品|同等規格", line):
            continue
        out.append(
            DocFinding(
                rule_id="D-026-指定品牌",
                article_no="26",
                article_title="技術規格不得不當限制競爭",
                severity="high",
                snippet=_ctx(text, m.start(), m.end()),
                advice="若指定廠牌/型號，須加註『或同等品』以符合 §26 第 3 項。",
            )
        )
    # 去重：同 snippet 只留一個
    seen = set()
    uniq = []
    for f in out:
        if f.snippet in seen:
            continue
        seen.add(f.snippet)
        uniq.append(f)
    return uniq


def check_restrictive_phrases(text: str) -> list[DocFinding]:
    """§26 / §37：限制競爭文字。"""
    patterns = [
        (r"限\s*(本國|國內|台灣|臺灣)\s*(製造|生產|廠商)", "限制原產地/國籍"),
        (r"不得\s*使用\s*[A-Za-z0-9一-龥]{2,10}\s*(廠牌|品牌)", "排他性品牌限制"),
        (r"僅限\s*[A-Za-z0-9一-龥]{2,10}\s*(原廠|授權)", "排他性授權限制"),
    ]
    out = []
    for pat, desc in patterns:
        for m in re.finditer(pat, text):
            if _is_negated_or_quoted(text, m.start()):
                continue
            out.append(
                DocFinding(
                    rule_id="D-026-排他文字",
                    article_no="26",
                    article_title="技術規格不得不當限制競爭",
                    severity="high",
                    snippet=_ctx(text, m.start(), m.end()),
                    advice=f"可能違反 §26：{desc}。如有正當理由，須於招標文件敘明。",
                )
            )
    return out


def _parse_cn_number(s: str) -> int | None:
    """簡易中文數字解析；只處理 RFI 常見組合（最多兩位數 + 萬/千/百，億/萬/千 單位另算）。"""
    s = s.strip().replace(",", "").replace("，", "")
    if s.isdigit():
        return int(s)
    # 支援「壹拾參」「二十」「三十五」等
    table = CN_DIGIT
    if not s or any(c not in table for c in s):
        return None
    if len(s) == 1:
        return table[s]
    # 處理「X十Y」或「十Y」或「X十」
    if "十" in s or "拾" in s:
        ten_char = "十" if "十" in s else "拾"
        a, _, b = s.partition(ten_char)
        head = table[a] if a else 1
        tail = table[b] if b else 0
        return head * 10 + tail
    return None


def _to_amount(num_str: str, unit: str) -> int | None:
    n = _parse_cn_number(num_str)
    if n is None:
        return None
    unit_factor = {
        "億": 100_000_000,
        "千萬": 10_000_000,
        "仟萬": 10_000_000,
        "百萬": 1_000_000,
        "萬": 10_000,
    }
    return n * unit_factor.get(unit, 0)


def check_qualification_threshold(text: str) -> list[DocFinding]:
    """§36/§37：廠商資格門檻不得逾越業務性質之需要。

    真正做數值比對：
    - 實收資本額：≥ 1 億 → high；≥ 3000 萬 → mid；其他不報
    - 近 N 年同類實績：≥ 20 件 → high；≥ 10 件 → mid
    """
    out = []
    # 實收資本額：抓「實收資本額 ... N 單位」
    cap_pat = re.compile(
        r"實收資本額[^\n。]{0,40}?"
        r"(\d{1,5}(?:[,，]\d{3})*|[一二三四五六七八九十壹貳參肆伍陸柒捌玖拾]{1,3})"
        r"\s*(億|仟萬|千萬|百萬|萬)"
    )
    for m in cap_pat.finditer(text):
        amount = _to_amount(m.group(1), m.group(2))
        if amount is None:
            continue
        # 排除「N 億以下／未滿／不超過」等上限敘述（多為中小企業定義之引文）
        tail = text[m.end():m.end() + 12]
        if re.match(r"\s*元?\s*(以下|未滿|不超過|不足|以內)", tail):
            continue
        # 排除「不得要求 / 依採購法 / 認定標準」等引文／否定
        if _is_negated_or_quoted(text, m.start()):
            continue
        if amount >= 100_000_000:
            sev = "high"
        elif amount >= 30_000_000:
            sev = "mid"
        else:
            continue  # 不報，避免誤判小額資本要求
        out.append(
            DocFinding(
                rule_id="D-036-資本額門檻",
                article_no="36",
                article_title="廠商資格之訂定",
                severity=sev,
                snippet=_ctx(text, m.start(), m.end()),
                advice=(
                    f"實收資本額門檻 ≈ NT${amount:,}，"
                    "請評估是否逾越業務性質需要（§37 不得不當限制競爭）。"
                ),
            )
        )
    # 近 N 年實績：M 件
    rec_pat = re.compile(
        r"近\s*[一二三四五六七八九十\d]+\s*年[^\n。]{0,30}?實績[^\n。]{0,20}?"
        r"(\d+|[一二三四五六七八九十壹貳參肆伍陸柒捌玖拾]+)\s*件"
    )
    for m in rec_pat.finditer(text):
        if _is_negated_or_quoted(text, m.start()):
            continue
        n = _parse_cn_number(m.group(1))
        if n is None:
            continue
        if n >= 20:
            sev = "high"
        elif n >= 10:
            sev = "mid"
        else:
            continue
        out.append(
            DocFinding(
                rule_id="D-036-實績門檻",
                article_no="36",
                article_title="廠商資格之訂定",
                severity=sev,
                snippet=_ctx(text, m.start(), m.end()),
                advice="實績件數門檻請確認與標的規模相稱，避免不當限制競爭（§37）。",
            )
        )
    return out


def check_short_waiting_period(text: str) -> list[DocFinding]:
    """§28：等標期不得短於法定下限（依招標方式/金額不同）。

    PoC：找出「等標期 X 日」，提示請核對法定最短期限。
    """
    out = []
    for m in re.finditer(r"等標期[^\n。]{0,20}?(\d+)\s*日", text):
        # 法定下限引文「等標期不得少於 N 日」不是本案實際期限
        if _is_lower_bound_quote(text, m.start(1)):
            continue
        days = int(m.group(1))
        sev = "high" if days < 7 else "mid" if days < 14 else "low"
        out.append(
            DocFinding(
                rule_id="D-028-等標期",
                article_no="28",
                article_title="等標期",
                severity=sev,
                snippet=_ctx(text, m.start(), m.end()),
                advice=f"等標期 {days} 日。請依招標方式與金額核對「招標期限標準」最短等標期。",
            )
        )
    return out


def check_restrictive_tender_basis(text: str) -> list[DocFinding]:
    """§22：限制性招標必須符合 §22 列舉款項。"""
    # 找出所有非否定的「限制性招標」出現位置
    occ = [m for m in re.finditer(r"限制性招標", text)
           if not _is_negated_or_quoted(text, m.start())]
    if not occ:
        return []
    # 是否有提到 §22 或其款項
    if re.search(r"第\s*二十二\s*條|第\s*22\s*條|§\s*22|採購法.{0,10}?第\s*22\s*條", text):
        return []
    m = occ[0]
    return [
        DocFinding(
            rule_id="D-022-依據缺漏",
            article_no="22",
            article_title="限制性招標適用條件",
            severity="mid",
            snippet=_ctx(text, m.start(), m.end()),
            advice="本文件採限制性招標但未引用 §22 各款依據，請補列適用款項。",
        )
    ]


def check_founding_years(text: str) -> list[DocFinding]:
    """§37：成立年限門檻過長可能不當限制競爭。

    偵測「公司設立／成立／登記 ... N 年以上」：
      ≥ 10 年 → high；≥ 5 年 → mid；< 5 年 → 不報。
    """
    out = []
    # 例：「公司設立登記年限應達 15 年以上」「成立滿 8 年」「登記滿三年」
    pat = re.compile(
        r"(設立(?:登記)?|成立|登記)[^\n。]{0,15}?"
        r"(\d{1,2}|[一二三四五六七八九十壹貳參肆伍陸柒捌玖拾]{1,3})"
        r"\s*年\s*(以上|以上之經驗|滿)?"
    )
    for m in pat.finditer(text):
        if _is_negated_or_quoted(text, m.start()):
            continue
        if _is_lower_bound_quote(text, m.start(2)):
            continue
        n = _parse_cn_number(m.group(2))
        if n is None:
            continue
        # 必須是「以上 / 滿」這種下限敘述才算門檻；單純「公司成立 3 年慶」不算
        tail_keyword = m.group(3) or text[m.end():m.end() + 4]
        if "以上" not in tail_keyword and "滿" not in tail_keyword:
            # 再放寬：往後 6 字內找
            tail2 = text[m.end():m.end() + 6]
            if "以上" not in tail2 and "滿" not in tail2:
                continue
        if n >= 10:
            sev = "high"
        elif n >= 5:
            sev = "mid"
        else:
            continue
        out.append(
            DocFinding(
                rule_id="D-037-成立年限",
                article_no="37",
                article_title="廠商資格不得不當限制競爭",
                severity=sev,
                snippet=_ctx(text, m.start(), m.end()),
                advice=(
                    f"要求公司成立 {n} 年以上，可能逾越業務需要（§37）。"
                    "成立年限通常無法直接反映履約能力，建議改以實績或人員經驗替代。"
                ),
            )
        )
    return out


def check_location_restriction(text: str) -> list[DocFinding]:
    """§37：限定公司登記地址 / 營業所所在縣市，多屬不當限制競爭。

    偵測「公司登記 / 設籍 / 所在地 ... 須位於 / 應設於 / 應為 + 縣市名」。
    """
    out = []
    # 縣市名：6 直轄市 + 主要縣市；不窮舉，採通用 [一-龥]{2,4}(市|縣)
    pat = re.compile(
        r"(公司\s*登記|公司\s*設籍|公司\s*所在地|公司\s*地址|登記地址|營業所|總公司)"
        r"[^\n。]{0,15}?"
        r"(須|應|限|限於|限定)\s*(位於|設於|設立於|登記於|為|在)?\s*"
        r"([一-龥]{2,4}(市|縣)(?:[、，及和或]\s*[一-龥]{2,4}(?:市|縣))?)"
    )
    for m in pat.finditer(text):
        if _is_negated_or_quoted(text, m.start()):
            continue
        out.append(
            DocFinding(
                rule_id="D-037-地理限定",
                article_no="37",
                article_title="廠商資格不得不當限制競爭",
                severity="high",
                snippet=_ctx(text, m.start(), m.end()),
                advice=(
                    f"限定公司登記/所在地於「{m.group(4)}」可能不當限制競爭（§37）。"
                    "除有正當地緣需要（如就近服務），應避免地理限定。"
                ),
            )
        )
    return out


def check_bid_bond_forfeiture(text: str) -> list[DocFinding]:
    """§31：押標金不予發還之事由必須限於法定款項（§31 第 2 項 5 款）。

    偵測：押標金 + 沒收/不予發還/繳庫 字眼，且文件未引用 §31。
    """
    if not re.search(r"押標金", text):
        return []
    forfeit_hits = list(
        re.finditer(r"押標金[^\n。]{0,40}?(不予發還|不發還|沒收|繳庫|歸機關所有)", text)
    )
    if not forfeit_hits:
        return []
    has_legal_basis = bool(
        re.search(r"第\s*三十一\s*條|第\s*31\s*條|§\s*31|採購法.{0,10}?第\s*31\s*條", text)
    )
    if has_legal_basis:
        return []
    m = forfeit_hits[0]
    return [
        DocFinding(
            rule_id="D-031-押標金沒收",
            article_no="31",
            article_title="押標金之發還及不予發還",
            severity="high",
            snippet=_ctx(text, m.start(), m.end()),
            advice="押標金不予發還事由須限於 §31 第 2 項各款；招標文件應明確引用法定事由。",
        )
    ]


def check_best_value_vague_criteria(text: str) -> list[DocFinding]:
    """§56 / 最有利標評選辦法：評選項目須具體、配分明確。

    偵測：文件提到最有利標 / 評選，但出現「綜合考量 / 綜合判斷 / 專業判斷」
    等籠統用語且無配分（無「配分」「%」「百分比」字樣）。
    """
    if not re.search(r"最有利標|評選", text):
        return []
    # 籠統用語必須出現在「最有利標 / 評選」鄰近上下文（±120 字）才算
    vague = None
    for m in re.finditer(r"綜合\s*(考量|判斷|評估)|專業判斷|綜合審酌", text):
        win = text[max(0, m.start() - 120):min(len(text), m.end() + 120)]
        if re.search(r"最有利標|評選", win):
            vague = m
            break
    if not vague:
        return []
    has_weights = bool(re.search(r"配分|百分比|權重|\d+\s*%|\d+\s*分", text))
    if has_weights:
        return []
    return [
        DocFinding(
            rule_id="D-056-評選籠統",
            article_no="56",
            article_title="最有利標決標",
            severity="mid",
            snippet=_ctx(text, vague.start(), vague.end()),
            advice="評選用語過於籠統（如「綜合考量」）且未見配分。應列具體評選項目與配分比例。",
        )
    ]


def check_short_performance_period(text: str) -> list[DocFinding]:
    """§70：履約期限應合理；過短可能變相限制競爭。

    PoC：履約期限 / 完成期限 < 7 日 → mid；< 14 日 → low。
    """
    out = []
    for m in re.finditer(r"(履約期限|完成期限|履約期間)[^\n。]{0,20}?(\d+)\s*日", text):
        if _is_lower_bound_quote(text, m.start(2)):
            continue
        days = int(m.group(2))
        if days >= 14:
            continue
        sev = "mid" if days < 7 else "low"
        out.append(
            DocFinding(
                rule_id="D-070-履約期限",
                article_no="70",
                article_title="履約管理",
                severity=sev,
                snippet=_ctx(text, m.start(), m.end()),
                advice=f"履約期限僅 {days} 日，請評估是否與標的規模相稱、是否變相限制競爭。",
            )
        )
    return out


CHECKS = [
    check_brand_without_equivalent,
    check_restrictive_phrases,
    check_qualification_threshold,
    check_short_waiting_period,
    check_restrictive_tender_basis,
    check_bid_bond_forfeiture,
    check_best_value_vague_criteria,
    check_short_performance_period,
    check_founding_years,
    check_location_restriction,
]


def run(text: str) -> list[DocFinding]:
    findings: list[DocFinding] = []
    for fn in CHECKS:
        findings.extend(fn(text))
    sev_order = {"high": 0, "mid": 1, "low": 2}
    findings.sort(key=lambda f: sev_order.get(f.severity, 3))
    return findings
